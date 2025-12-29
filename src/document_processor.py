import os
import tempfile
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import markdown
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """处理各种格式的文档"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
        )
    
    def load_document(self, file_path: str) -> Optional[List[LangchainDocument]]:
        """加载并处理单个文档"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None
        
        try:
            if file_path.suffix.lower() == ".pdf":
                return self._load_pdf(file_path)
            elif file_path.suffix.lower() == ".docx":
                return self._load_docx(file_path)
            elif file_path.suffix.lower() in [".txt", ".md"]:
                return self._load_text(file_path)
            else:
                logger.error(f"不支持的文件格式: {file_path.suffix}")
                return None
        except Exception as e:
            logger.error(f"处理文件时出错 {file_path}: {str(e)}")
            return None
    
    def _load_pdf(self, file_path: Path) -> List[LangchainDocument]:
        """加载PDF文件"""
        documents = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    metadata = {
                        "source": str(file_path),
                        "page": page_num + 1,
                        "type": "pdf"
                    }
                    documents.append(LangchainDocument(
                        page_content=text,
                        metadata=metadata
                    ))
        return documents
    
    def _load_docx(self, file_path: Path) -> List[LangchainDocument]:
        """加载Word文档"""
        documents = []
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text = "\n".join(full_text)
        if text.strip():
            metadata = {
                "source": str(file_path),
                "type": "docx"
            }
            documents.append(LangchainDocument(
                page_content=text,
                metadata=metadata
            ))
        
        return documents
    
    def _load_text(self, file_path: Path) -> List[LangchainDocument]:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        metadata = {
            "source": str(file_path),
            "type": file_path.suffix[1:]
        }
        
        return [LangchainDocument(
            page_content=text,
            metadata=metadata
        )]
    
    def chunk_documents(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """将文档分割成块"""
        return self.text_splitter.split_documents(documents)
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """处理目录中的所有文档"""
        all_documents = []
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            logger.error(f"目录不存在: {directory_path}")
            return all_documents
        
        for file_path in directory.rglob("*"):
            if file_path.suffix.lower() in [".pdf", ".docx", ".txt", ".md"]:
                documents = self.load_document(file_path)
                if documents:
                    all_documents.extend(documents)
        
        return self.chunk_documents(all_documents)
