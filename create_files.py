import os
from pathlib import Path

# 定义文件内容
files_content = {
    # 文档处理器
    "src/document_processor.py": '''import os
import tempfile
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """处理各种格式的文档"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\\n\\n", "\\n", "。", "！", "？", "；", "，", " ", ""]
        )
    
    def load_document(self, file_path: str) -> Optional[List[LangchainDocument]]:
        """加载并处理单个文档"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None
        
        try:
            if file_path.suffix.lower() == ".pdf":
                return self._load_pdf(file_path)
            elif file_path.suffix.lower() == ".docx":
                return self._load_docx(file_path)
            elif file_path.suffix.lower() in [".txt", ".md"]:
                return self._load_text(file_path)
            else:
                logger.error(f"不支持的文件格式: {file_path.suffix}")
                return None
        except Exception as e:
            logger.error(f"处理文件时出错 {file_path}: {str(e)}")
            return None
    
    def _load_pdf(self, file_path: Path) -> List[LangchainDocument]:
        """加载PDF文件"""
        documents = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    metadata = {
                        "source": str(file_path),
                        "page": page_num + 1,
                        "type": "pdf"
                    }
                    documents.append(LangchainDocument(
                        page_content=text,
                        metadata=metadata
                    ))
        return documents
    
    def _load_docx(self, file_path: Path) -> List[LangchainDocument]:
        """加载Word文档"""
        documents = []
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text = "\\n".join(full_text)
        if text.strip():
            metadata = {
                "source": str(file_path),
                "type": "docx"
            }
            documents.append(LangchainDocument(
                page_content=text,
                metadata=metadata
            ))
        
        return documents
    
    def _load_text(self, file_path: Path) -> List[LangchainDocument]:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        metadata = {
            "source": str(file_path),
            "type": file_path.suffix[1:]
        }
        
        return [LangchainDocument(
            page_content=text,
            metadata=metadata
        )]
    
    def chunk_documents(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """将文档分割成块"""
        return self.text_splitter.split_documents(documents)
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """处理目录中的所有文档"""
        all_documents = []
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            logger.error(f"目录不存在: {directory_path}")
            return all_documents
        
        for file_path in directory.rglob("*"):
            if file_path.suffix.lower() in [".pdf", ".docx", ".txt", ".md"]:
                documents = self.load_document(file_path)
                if documents:
                    all_documents.extend(documents)
        
        return self.chunk_documents(all_documents)
''',
    
    # 向量存储
    "src/vector_store.py": '''import os
from typing import List, Optional, Dict, Any
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.schema import Document as LangchainDocument
import chromadb
from chromadb.config import Settings as ChromaSettings
from config.settings import settings
import logging

logger = logging.getLogger(__name__)

class VectorStoreManager:
    """管理向量数据库"""
    
    def __init__(self):
        self.embeddings = self._create_embeddings()
        self.client_settings = ChromaSettings(
            chroma_db_impl="duckdb+parquet",
            persist_directory=settings.CHROMA_PERSIST_DIR,
            anonymized_telemetry=False
        )
        
    def _create_embeddings(self):
        """创建嵌入模型"""
        return OpenAIEmbeddings(
            model=settings.EMBEDDING_MODEL,
            openai_api_key=settings.DEEPSEEK_API_KEY,
            openai_api_base=settings.DEEPSEEK_API_BASE
        )
    
    def create_vector_store(self, 
                           documents: List[LangchainDocument],
                           collection_name: str = "research_docs") -> Chroma:
        """创建新的向量存储"""
        if os.path.exists(settings.CHROMA_PERSIST_DIR):
            import shutil
            shutil.rmtree(settings.CHROMA_PERSIST_DIR)
        
        vector_store = Chroma.from_documents(
            documents=documents,
            embedding=self.embeddings,
            persist_directory=settings.CHROMA_PERSIST_DIR,
            collection_name=collection_name,
            client_settings=self.client_settings
        )
        
        vector_store.persist()
        logger.info(f"向量存储已创建，包含 {len(documents)} 个文档块")
        return vector_store
    
    def load_vector_store(self, collection_name: str = "research_docs") -> Optional[Chroma]:
        """加载现有的向量存储"""
        try:
            vector_store = Chroma(
                persist_directory=settings.CHROMA_PERSIST_DIR,
                embedding_function=self.embeddings,
                collection_name=collection_name,
                client_settings=self.client_settings
            )
            
            collection = vector_store._collection
            if collection.count() == 0:
                logger.warning("向量存储为空")
                return None
            
            logger.info(f"向量存储已加载，包含 {collection.count()} 个文档")
            return vector_store
        except Exception as e:
            logger.error(f"加载向量存储时出错: {str(e)}")
            return None
    
    def search_similar(self, 
                      query: str, 
                      vector_store: Chroma, 
                      k: int = 5) -> List[Dict[str, Any]]:
        """搜索相似文档"""
        results = vector_store.similarity_search_with_score(query, k=k)
        
        formatted_results = []
        for doc, score in results:
            formatted_results.append({
                "content": doc.page_content,
                "metadata": doc.metadata,
                "score": score
            })
        
        return formatted_results
    
    def delete_collection(self, collection_name: str = "research_docs"):
        """删除向量存储"""
        try:
            client = chromadb.Client(self.client_settings)
            client.delete_collection(collection_name)
            logger.info(f"集合 '{collection_name}' 已删除")
        except Exception as e:
            logger.error(f"删除集合时出错: {str(e)}")
''',
    
    # 工具类
    "src/tools.py": '''from typing import Type, Optional, Dict, Any
from langchain.tools import BaseTool, Tool
from langchain.callbacks.manager import CallbackManagerForToolRun
from langchain.utilities import SerpAPIWrapper
import requests
from datetime import datetime
import json
import logging

logger = logging.getLogger(__name__)

class SearchTool(BaseTool):
    """搜索工具"""
    name: str = "search"
    description: str = """
    用于搜索最新信息。当需要最新数据、新闻或特定事实时使用此工具。
    输入应为搜索查询字符串。
    """
    
    def __init__(self, api_key: Optional[str] = None):
        super().__init__()
        self.has_api = api_key is not None
        if self.has_api:
            self.search_wrapper = SerpAPIWrapper(serpapi_api_key=api_key)
    
    def _run(self, query: str, run_manager: Optional[CallbackManagerForToolRun] = None) -> str:
        try:
            if self.has_api:
                return self.search_wrapper.run(query)
            else:
                return f"搜索查询: {query}\\n\\n注意：搜索功能需要配置API密钥。请使用本地知识库或提供具体问题。"
        except Exception as e:
            logger.error(f"搜索时出错: {str(e)}")
            return f"搜索失败: {str(e)}"

class CalculatorTool(BaseTool):
    """计算器工具"""
    name: str = "calculator"
    description: str = """
    用于执行数学计算。当需要计算数字、百分比、统计量等时使用此工具。
    输入应为数学表达式字符串。
    """
    
    def _run(self, expression: str, run_manager: Optional[CallbackManagerForToolRun] = None) -> str:
        try:
            allowed_chars = set("0123456789+-*/(). ")
            if not all(c in allowed_chars for c in expression):
                return "错误：表达式中包含非法字符"
            
            result = eval(expression, {"__builtins__": {}}, {})
            return f"计算结果: {expression} = {result}"
        except Exception as e:
            return f"计算错误: {str(e)}"

class CurrentDateTimeTool(BaseTool):
    """获取当前日期时间工具"""
    name: str = "get_current_datetime"
    description: str = """
    获取当前的日期和时间。当需要知道当前时间或计算时间相关问题时使用此工具。
    输入可以是任何字符串，通常为空。
    """
    
    def _run(self, query: str = "", run_manager: Optional[CallbackManagerForToolRun] = None) -> str:
        now = datetime.now()
        return f"当前日期和时间: {now.strftime('%Y年%m月%d日 %H:%M:%S')}"

class KnowledgeBaseQueryTool(BaseTool):
    """知识库查询工具"""
    name: str = "query_knowledge_base"
    description: str = """
    从本地知识库中检索相关信息。当需要参考已上传的文档时使用此工具。
    输入应为要查询的问题或关键词。
    """
    
    def __init__(self, vector_store_manager, **kwargs):
        super().__init__(**kwargs)
        self.vector_store_manager = vector_store_manager
        self.vector_store = None
    
    def _run(self, query: str, run_manager: Optional[CallbackManagerForToolRun] = None) -> str:
        try:
            if self.vector_store is None:
                self.vector_store = self.vector_store_manager.load_vector_store()
            
            if self.vector_store is None:
                return "知识库未加载或为空。请先上传文档。"
            
            results = self.vector_store_manager.search_similar(query, self.vector_store, k=3)
            
            if not results:
                return "未找到相关信息。"
            
            response = "从知识库中找到以下相关信息：\\n\\n"
            for i, result in enumerate(results, 1):
                response += f"{i}. 来源: {result['metadata'].get('source', '未知')}\\n"
                if 'page' in result['metadata']:
                    response += f"   页码: {result['metadata']['page']}\\n"
                response += f"   内容: {result['content'][:200]}...\\n"
                response += f"   相关性: {result['score']:.3f}\\n\\n"
            
            return response
        except Exception as e:
            logger.error(f"查询知识库时出错: {str(e)}")
            return f"查询知识库时出错: {str(e)}"

def create_tools(vector_store_manager=None, search_api_key=None):
    """创建所有工具"""
    tools = []
    
    tools.append(SearchTool(api_key=search_api_key))
    tools.append(CalculatorTool())
    tools.append(CurrentDateTimeTool())
    
    if vector_store_manager:
        tools.append(KnowledgeBaseQueryTool(
            vector_store_manager=vector_store_manager,
            name="query_knowledge_base",
            description="从已上传的文档中检索相关信息"
        ))
    
    return tools
''',
    
    # 代理
    "src/agent.py": '''from typing import List, Dict, Any, Optional
from langchain.agents import AgentExecutor, create_react_agent
from langchain.prompts import PromptTemplate
from langchain.tools import BaseTool
from langchain.memory import ConversationBufferMemory
from langchain.callbacks import StreamingStdOutCallbackHandler
from config.settings import settings
from langchain_community.chat_models import ChatOpenAI
import logging

logger = logging.getLogger(__name__)

class ResearchAssistant:
    """研究助手代理"""
    
    def __init__(self, tools: List[BaseTool], model_name: str = None):
        self.tools = tools
        self.model_name = model_name or settings.DEEPSEEK_MODEL
        
        self.llm = self._create_llm()
        
        self.memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True,
            output_key="output"
        )
        
        self.prompt = self._create_prompt()
        
        self.agent = create_react_agent(
            llm=self.llm,
            tools=self.tools,
            prompt=self.prompt
        )
        
        self.agent_executor = AgentExecutor(
            agent=self.agent,
            tools=self.tools,
            memory=self.memory,
            verbose=True,
            handle_parsing_errors=True,
            max_iterations=settings.MAX_ITERATIONS,
            return_intermediate_steps=True
        )
    
    def _create_llm(self):
        """创建DeepSeek LLM实例"""
        return ChatOpenAI(
            model=self.model_name,
            temperature=settings.TEMPERATURE,
            openai_api_key=settings.DEEPSEEK_API_KEY,
            openai_api_base=settings.DEEPSEEK_API_BASE,
            streaming=False,
            callbacks=[StreamingStdOutCallbackHandler()]
        )
    
    def _create_prompt(self) -> PromptTemplate:
        """创建代理提示模板"""
        template = """你是一个专业的研究助手，拥有广泛的知识和强大的研究能力。
        
        你可以使用以下工具：
        {tools}
        
        使用以下格式：
        问题：用户提出的问题
        思考：你需要思考如何解决问题
        行动：要采取的行动，应该是[{tool_names}]中的一个
        行动输入：行动的输入
        观察：行动的结果
        ...（这个思考/行动/行动输入/观察可以重复多次）
        思考：我现在知道了最终答案
        最终答案：对原始问题的最终回答
        
        请用中文回答，并且回答应该详细、准确、有引用来源。
        
        如果你需要引用来源，请明确指出信息来源。
        
        之前的对话：
        {chat_history}
        
        开始！
        
        问题：{input}
        思考：{agent_scratchpad}"""
        
        return PromptTemplate.from_template(template)
    
    def query(self, question: str) -> Dict[str, Any]:
        """向助手提问"""
        try:
            result = self.agent_executor.invoke({
                "input": question
            })
            
            return {
                "answer": result.get("output", "没有获取到回答"),
                "intermediate_steps": result.get("intermediate_steps", []),
                "success": True
            }
        except Exception as e:
            logger.error(f"执行查询时出错: {str(e)}")
            return {
                "answer": f"抱歉，处理请求时出现错误: {str(e)}",
                "success": False
            }
    
    def reset_memory(self):
        """重置对话记忆"""
        self.memory.clear()
        logger.info("对话记忆已重置")
    
    def get_chat_history(self) -> List[Dict[str, str]]:
        """获取对话历史"""
        memory_variables = self.memory.load_memory_variables({})
        chat_history = memory_variables.get("chat_history", [])
        
        formatted_history = []
        for i in range(0, len(chat_history), 2):
            if i + 1 < len(chat_history):
                formatted_history.append({
                    "human": chat_history[i].content,
                    "ai": chat_history[i + 1].content
                })
        
        return formatted_history
''',
    
    # 研究报告生成器
    "src/research_writer.py": '''from typing import List, Dict, Any, Optional
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from datetime import datetime
import markdown
import pdfkit
import os
from pathlib import Path
from config.settings import settings
import logging

logger = logging.getLogger(__name__)

class ResearchWriter:
    """研究报告生成器"""
    
    def __init__(self, llm):
        self.llm = llm
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(self, 
                       topic: str, 
                       research_data: Dict[str, Any],
                       format: str = "markdown") -> str:
        """生成研究报告"""
        
        template = """# 研究报告：{topic}
        
生成时间：{timestamp}
        
## 摘要
        
{summary}
        
## 研究背景
        
{background}
        
## 关键发现
        
{findings}
        
## 详细分析
        
{analysis}
        
## 结论
        
{conclusion}
        
## 参考文献
        
{references}
        
## 附录：研究过程
        
{process}
        """
        
        report_data = {
            "topic": topic,
            "timestamp": datetime.now().strftime("%Y年%m月%d日 %H:%M:%S"),
            "summary": research_data.get("summary", ""),
            "background": research_data.get("background", ""),
            "findings": self._format_findings(research_data.get("findings", [])),
            "analysis": research_data.get("analysis", ""),
            "conclusion": research_data.get("conclusion", ""),
            "references": self._format_references(research_data.get("references", [])),
            "process": research_data.get("process", "")
        }
        
        report = template.format(**report_data)
        
        filename = self._save_report(topic, report, format)
        
        return filename
    
    def _format_findings(self, findings: List[str]) -> str:
        """格式化关键发现"""
        if not findings:
            return "暂无关键发现"
        
        formatted = ""
        for i, finding in enumerate(findings, 1):
            formatted += f"{i}. {finding}\\n"
        
        return formatted
    
    def _format_references(self, references: List[Dict[str, str]]) -> str:
        """格式化参考文献"""
        if not references:
            return "暂无参考文献"
        
        formatted = ""
        for i, ref in enumerate(references, 1):
            formatted += f"{i}. {ref.get('title', '无标题')}\\n"
            if ref.get('author'):
                formatted += f"   作者: {ref['author']}\\n"
            if ref.get('source'):
                formatted += f"   来源: {ref['source']}\\n"
            if ref.get('date'):
                formatted += f"   日期: {ref['date']}\\n"
            formatted += "\\n"
        
        return formatted
    
    def _save_report(self, topic: str, report: str, format: str) -> str:
        """保存报告到文件"""
        safe_topic = "".join(c for c in topic if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_topic = safe_topic.replace(' ', '_')[:50]
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format.lower() == "markdown":
            filename = self.output_dir / f"{safe_topic}_{timestamp}.md"
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(report)
            logger.info(f"Markdown报告已保存: {filename}")
            
            html_filename = self.output_dir / f"{safe_topic}_{timestamp}.html"
            html_content = markdown.markdown(report, extensions=['tables', 'fenced_code'])
            with open(html_filename, 'w', encoding='utf-8') as f:
                f.write(f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>研究报告: {topic}</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
                        h1 {{ color: #333; border-bottom: 2px solid #eee; }}
                        h2 {{ color: #555; }}
                        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
                        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """)
            logger.info(f"HTML报告已保存: {html_filename}")
            
        elif format.lower() == "pdf":
            filename = self.output_dir / f"{safe_topic}_{timestamp}.pdf"
            try:
                pdfkit.from_string(report, str(filename))
                logger.info(f"PDF报告已保存: {filename}")
            except Exception as e:
                logger.error(f"生成PDF时出错: {str(e)}")
                filename = self._save_report(topic, report, "markdown")
        else:
            filename = self.output_dir / f"{safe_topic}_{timestamp}.txt"
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(report)
            logger.info(f"文本报告已保存: {filename}")
        
        return str(filename)
    
    def generate_summary_chain(self) -> LLMChain:
        """创建摘要生成链"""
        prompt = PromptTemplate(
            input_variables=["content"],
            template="请为以下内容生成一个简洁的摘要：\\n\\n{content}\\n\\n摘要："
        )
        return LLMChain(llm=self.llm, prompt=prompt)
''',
    
    # 工具函数
    "src/utils.py": '''import logging
import sys
from pathlib import Path
from typing import Optional
import json

def setup_logging(level: int = logging.INFO, log_file: Optional[str] = None):
    """设置日志配置"""
    log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    
    handlers = [logging.StreamHandler(sys.stdout)]
    
    if log_file:
        Path(log_file).parent.mkdir(parents=True, exist_ok=True)
        handlers.append(logging.FileHandler(log_file, encoding='utf-8'))
    
    logging.basicConfig(
        level=level,
        format=log_format,
        handlers=handlers
    )

def save_conversation(conversation: list, filepath: str):
    """保存对话记录"""
    Path(filepath).parent.mkdir(parents=True, exist_ok=True)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(conversation, f, ensure_ascii=False, indent=2)

def load_conversation(filepath: str) -> list:
    """加载对话记录"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return []

def format_document_for_display(document: dict, max_length: int = 300) -> str:
    """格式化文档用于显示"""
    content = document.get('content', '')
    if len(content) > max_length:
        content = content[:max_length] + '...'
    
    metadata = document.get('metadata', {})
    source = metadata.get('source', '未知来源')
    page = metadata.get('page', '')
    
    result = f"来源: {source}"
    if page:
        result += f" (页 {page})"
    result += f"\\n内容: {content}\\n"
    
    return result
''',
    
    # 主应用程序
    "src/main.py": '''import sys
import os
from pathlib import Path
sys.path.append(str(Path(__file__).parent.parent))

from typing import Optional, Dict, Any
import logging
from colorama import init, Fore, Style

from config.settings import settings
from src.document_processor import DocumentProcessor
from src.vector_store import VectorStoreManager
from src.tools import create_tools
from src.agent import ResearchAssistant
from src.research_writer import ResearchWriter
from src.utils import setup_logging

init(autoreset=True)

class IntelligentResearchAssistant:
    """智能研究助手主类"""
    
    def __init__(self):
        setup_logging()
        self.logger = logging.getLogger(__name__)
        
        if not settings.DEEPSEEK_API_KEY:
            print(Fore.YELLOW + "警告: 未设置DeepSeek API密钥。请设置DEEPSEEK_API_KEY环境变量。")
        
        self.document_processor = DocumentProcessor()
        self.vector_store_manager = VectorStoreManager()
        self.vector_store = None
        
        self.tools = create_tools(
            vector_store_manager=self.vector_store_manager,
            search_api_key=os.getenv("SERPAPI_API_KEY")
        )
        
        self.assistant = ResearchAssistant(self.tools)
        self.writer = ResearchWriter(self.assistant.llm)
        
        self.logger.info("智能研究助手初始化完成")
    
    def load_documents(self, path: str):
        """加载文档到知识库"""
        print(Fore.CYAN + f"正在加载文档: {path}")
        
        if os.path.isfile(path):
            documents = self.document_processor.load_document(path)
            if documents:
                documents = self.document_processor.chunk_documents(documents)
        elif os.path.isdir(path):
            documents = self.document_processor.process_directory(path)
        else:
            print(Fore.RED + f"路径不存在: {path}")
            return False
        
        if not documents:
            print(Fore.RED + "未找到可处理的文档")
            return False
        
        print(Fore.GREEN + f"成功加载 {len(documents)} 个文档块")
        
        self.vector_store = self.vector_store_manager.create_vector_store(documents)
        return True
    
    def ask_question(self, question: str) -> Dict[str, Any]:
        """提问并获取回答"""
        print(Fore.YELLOW + f"\\n问题: {question}")
        print(Fore.CYAN + "思考中...\\n")
        
        result = self.assistant.query(question)
        
        if result["success"]:
            print(Fore.GREEN + "\\n" + "="*50)
            print(Fore.GREEN + "回答:")
            print(Fore.GREEN + "="*50)
            print(Fore.WHITE + result["answer"])
            print(Fore.GREEN + "="*50)
        else:
            print(Fore.RED + f"错误: {result['answer']}")
        
        return result
    
    def interactive_chat(self):
        """交互式聊天模式"""
        print(Fore.CYAN + "="*60)
        print(Fore.CYAN + "智能研究助手已启动 (输入 '退出' 或 'quit' 结束)")
        print(Fore.CYAN + "="*60)
        
        while True:
            try:
                question = input(Fore.YELLOW + "\\n您的问题: " + Style.RESET_ALL)
                
                if question.lower() in ['退出', 'quit', 'exit', 'q']:
                    print(Fore.CYAN + "再见！")
                    break
                
                if question.strip():
                    self.ask_question(question)
            except KeyboardInterrupt:
                print(Fore.CYAN + "\\n\\n再见！")
                break
            except Exception as e:
                print(Fore.RED + f"发生错误: {str(e)}")
    
    def generate_research_report(self, topic: str, questions: list):
        """生成研究报告"""
        print(Fore.CYAN + f"开始生成研究报告: {topic}")
        
        research_data = {
            "summary": "",
            "background": "",
            "findings": [],
            "analysis": "",
            "conclusion": "",
            "references": [],
            "process": ""
        }
        
        all_answers = []
        for i, question in enumerate(questions, 1):
            print(Fore.YELLOW + f"\\n[{i}/{len(questions)}] 研究问题: {question}")
            result = self.ask_question(question)
            
            if result["success"]:
                all_answers.append({
                    "question": question,
                    "answer": result["answer"],
                    "steps": result.get("intermediate_steps", [])
                })
        
        research_data["process"] = self._format_research_process(all_answers)
        
        summary_chain = self.writer.generate_summary_chain()
        combined_content = "\\n\\n".join([a["answer"] for a in all_answers])
        research_data["summary"] = summary_chain.run(content=combined_content)
        
        report_file = self.writer.generate_report(topic, research_data, "markdown")
        
        print(Fore.GREEN + f"\\n研究报告已生成: {report_file}")
        return report_file
    
    def _format_research_process(self, answers: list) -> str:
        """格式化研究过程"""
        process = "## 研究过程记录\\n\\n"
        
        for i, answer in enumerate(answers, 1):
            process += f"### 问题 {i}: {answer['question']}\\n\\n"
            process += f"**回答**:\\n{answer['answer']}\\n\\n"
            
            if answer.get('steps'):
                process += "**推理步骤**:\\n"
                for step in answer['steps']:
                    if isinstance(step, tuple) and len(step) >= 2:
                        action, observation = step[0], step[1]
                        process += f"- 行动: {action}\\n"
                        process += f"  结果: {observation[:200]}...\\n\\n"
        
        return process
    
    def show_capabilities(self):
        """显示助手功能"""
        print(Fore.CYAN + "="*60)
        print(Fore.CYAN + "智能研究助手功能")
        print(Fore.CYAN + "="*60)
        print(Fore.YELLOW + "1. 文档处理")
        print(Fore.WHITE + "   - 支持PDF、DOCX、TXT、MD格式")
        print(Fore.WHITE + "   - 自动分块和嵌入")
        print(Fore.YELLOW + "2. 知识库检索")
        print(Fore.WHITE + "   - 基于向量的语义搜索")
        print(Fore.WHITE + "   - 相关文档推荐")
        print(Fore.YELLOW + "3. 研究能力")
        print(Fore.WHITE + "   - 多步骤推理")
        print(Fore.WHITE + "   - 工具使用（搜索、计算等）")
        print(Fore.YELLOW + "4. 报告生成")
        print(Fore.WHITE + "   - 自动生成研究报告")
        print(Fore.WHITE + "   - 支持多种格式输出")
        print(Fore.CYAN + "="*60)

def main():
    """主函数"""
    print(Fore.CYAN + """
    ╔══════════════════════════════════════════════════════════╗
    ║           基于LangChain与DeepSeek的智能研究助手           ║
    ╚══════════════════════════════════════════════════════════╝
    """)
    
    assistant = IntelligentResearchAssistant()
    
    while True:
        print(Fore.CYAN + "\\n" + "="*60)
        print(Fore.CYAN + "主菜单")
        print(Fore.CYAN + "="*60)
        print(Fore.YELLOW + "1. 交互式聊天")
        print(Fore.YELLOW + "2. 加载文档到知识库")
        print(Fore.YELLOW + "3. 生成研究报告")
        print(Fore.YELLOW + "4. 显示功能")
        print(Fore.YELLOW + "5. 退出")
        print(Fore.CYAN + "="*60)
        
        choice = input(Fore.GREEN + "请选择 (1-5): " + Style.RESET_ALL)
        
        if choice == "1":
            assistant.interactive_chat()
        elif choice == "2":
            path = input(Fore.GREEN + "请输入文档路径或目录: " + Style.RESET_ALL)
            assistant.load_documents(path)
        elif choice == "3":
            topic = input(Fore.GREEN + "请输入研究主题: " + Style.RESET_ALL)
            print(Fore.GREEN + "请输入研究问题（每行一个问题，空行结束）:")
            questions = []
            while True:
                q = input(Fore.WHITE + "> " + Style.RESET_ALL)
                if not q.strip():
                    break
                questions.append(q)
            
            if questions:
                assistant.generate_research_report(topic, questions)
            else:
                print(Fore.RED + "未输入任何问题")
        elif choice == "4":
            assistant.show_capabilities()
        elif choice == "5":
            print(Fore.CYAN + "再见！")
            break
        else:
            print(Fore.RED + "无效选择，请重新输入")

if __name__ == "__main__":
    main()
''',
    
    # 根目录的main.py
    "main.py": '''#!/usr/bin/env python3
"""
智能研究助手入口点
"""
import sys
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.append(str(project_root))

from src.main import main

if __name__ == "__main__":
    main()
''',
    
    # 环境变量示例
    ".env.example": '''# DeepSeek API配置
DEEPSEEK_API_KEY=your_deepseek_api_key_here
DEEPSEEK_API_BASE=https://api.deepseek.com
DEEPSEEK_MODEL=deepseek-chat

# 可选：搜索API（如SerpAPI）
# SERPAPI_API_KEY=your_serpapi_api_key_here

# 向量数据库配置
EMBEDDING_MODEL=text-embedding-ada-002

# 应用配置
MAX_ITERATIONS=10
TEMPERATURE=0.1
''',
    
    # 依赖文件
    "requirements.txt": '''langchain==0.1.0
langchain-community==0.0.10
langchain-openai==0.0.5
openai==1.3.0
chromadb==0.4.18
pypdf2==3.0.1
python-docx==1.1.0
markdown==3.5.1
pydantic-settings==2.1.0
python-dotenv==1.0.0
colorama==0.4.6
requests==2.31.0
pdfkit==1.0.0
# 可选：用于搜索功能
# google-search-results==2.4.2
''',
}

# 创建所有文件
for filepath, content in files_content.items():
    file_path = Path(filepath)
    file_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"已创建: {filepath}")

print("\\n所有文件创建完成！")
print("现在请运行以下命令：")
print("1. cp .env.example .env")
print("2. 编辑 .env 文件，添加您的DeepSeek API密钥")
print("3. pip install -r requirements.txt")
print("4. python main.py")